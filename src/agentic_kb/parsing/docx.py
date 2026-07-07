"""DOCX parser backed by python-docx with optional Qwen VL image description."""

from collections.abc import Callable
from io import BytesIO

from agentic_kb.parsing.base import (
    Parser,
    ParserLimitError,
    ParserReadError,
    UnsupportedContentTypeError,
    raise_parser_read_error,
)
from agentic_kb.parsing.schemas import ParsedDocument, ParsedElement, ParsedSection

ImageDescriber = Callable[[bytes, str], str]


def _default_image_describer() -> ImageDescriber:
    """Return a describer that lazily initialises Qwen only when called."""
    _qwen: object | None = None

    def describe(image_bytes: bytes, prompt: str) -> str:
        nonlocal _qwen
        if _qwen is None:
            from agentic_kb.providers import Qwen
            _qwen = Qwen()
        return _qwen.describe_image(image_bytes, prompt)  # type: ignore[union-attr]

    return describe


class DocxParser(Parser):
    """Parse DOCX paragraphs and heading styles via python-docx.

    Embedded images are described via an image_describer (default: Qwen VL).
    Descriptions are inserted as text elements so downstream RAG can match them.
    """

    supported_content_types = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    def __init__(self, *, image_describer: ImageDescriber | None = None) -> None:
        super().__init__()
        self._image_describer = image_describer or _default_image_describer()

    def parse_text(
        self,
        text: str,
        *,
        source_uri: str,
        content_type: str,
    ) -> ParsedDocument:
        raise TypeError("DocxParser requires bytes content")

    def parse(
        self,
        content: bytes | str,
        *,
        source_uri: str,
        content_type: str,
    ) -> ParsedDocument:
        if isinstance(content, str):
            content = content.encode("utf-8")
        self._ensure_supported(content_type)
        self._validate_content_size(content)
        normalized_content_type = content_type.split(";", 1)[0].strip().lower()

        try:
            from docx import Document

            doc = Document(BytesIO(content))
            elements, image_tasks = _extract_elements(doc)
            elements = _describe_and_insert_images(
                doc, elements, image_tasks, self._image_describer,
            )
            sections = _sections_from_elements(elements)
        except (ParserLimitError, ParserReadError, UnsupportedContentTypeError):
            raise
        except Exception as error:
            raise_parser_read_error(source_uri, normalized_content_type, error)

        document = ParsedDocument(
            source_uri=source_uri,
            content_type=normalized_content_type,
            sections=sections,
            elements=elements,
            metadata={"paragraph_count": len(elements)},
        )
        self._validate_document_limits(document)
        return document

    def _ensure_supported(self, content_type: str) -> None:
        normalized = content_type.split(";", 1)[0].strip().lower()
        if normalized not in self.supported_content_types:
            raise UnsupportedContentTypeError(f"DocxParser does not support {content_type}")


def _extract_paragraphs(content: bytes) -> list[str]:
    """Return paragraph text only; kept as a small compatibility helper."""
    from docx import Document

    doc = Document(BytesIO(content))
    return [para.text.strip() for para in doc.paragraphs if para.text.strip()]


# ---------------------------------------------------------------------------
# Step 2 — paragraph / heading / image-ref extraction via python-docx
# ---------------------------------------------------------------------------

def _extract_elements(doc: "Document") -> tuple[list[ParsedElement], list[dict]]:
    """Extract paragraphs, headings and image references from a python-docx Document.

    Returns (elements, image_tasks).  Each image_task is:
        {"paragraph_index": int, "rel_id": str}
    """
    elements: list[ParsedElement] = []
    image_tasks: list[dict] = []
    heading_stack: list[str] = []
    paragraph_index = 0

    # Namespace for finding blip embeds in inline shapes.
    DML = "{http://schemas.openxmlformats.org/drawingml/2006/main}"
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    for para in doc.paragraphs:
        text = para.text.strip()

        # Detect images in this paragraph via XML children.
        has_image = False
        for run in para._element.iterchildren():
            # <w:drawing> → <wp:inline> → <a:graphic> → <a:graphicData> → <pic:pic> → <pic:blipFill> → <a:blip>
            for blip in run.iter(f"{DML}blip"):
                embed = blip.attrib.get(f"{{{R_NS}}}embed")
                if embed:
                    image_tasks.append({"paragraph_index": paragraph_index, "rel_id": embed})
                    has_image = True

        # Image-only paragraph: placeholder so the description lands in the right spot.
        if not text and has_image:
            elements.append(
                ParsedElement(
                    index=len(elements),
                    kind="paragraph",
                    text="",
                    section_path=tuple(heading_stack),
                )
            )
            paragraph_index += 1
            continue

        if not text:
            continue

        heading_level = _heading_level_from_style(para.style.name if para.style else None)
        if heading_level is not None:
            heading_stack = heading_stack[: heading_level - 1]
            heading_stack.append(text)
            elements.append(
                ParsedElement(
                    index=len(elements),
                    kind="heading",
                    text=text,
                    section_path=tuple(heading_stack),
                    metadata={"level": heading_level},
                )
            )
            paragraph_index += 1
            continue

        elements.append(
            ParsedElement(
                index=len(elements),
                kind="paragraph",
                text=text,
                section_path=tuple(heading_stack),
            )
        )
        paragraph_index += 1

    # # Extract tables — python-docx stores them separately from paragraphs.
    # for table in doc.tables:
    #     rows: list[str] = []
    #     for row in table.rows:
    #         cells = [cell.text.strip() for cell in row.cells]
    #         rows.append(" | ".join(cells))
    #     table_text = "\n".join(rows)
    #     if table_text.strip():
    #         elements.append(
    #             ParsedElement(
    #                 index=len(elements),
    #                 kind="table",
    #                 text=table_text,
    #                 section_path=tuple(heading_stack),
    #             )
    #         )

    # # Extract headers and footers from all sections.
    # for section in doc.sections:
    #     for header in (section.header, section.first_page_header, section.even_page_header):
    #         if header is None:
    #             continue
    #         text = "\n".join(p.text.strip() for p in header.paragraphs if p.text.strip())
    #         if text:
    #             elements.append(
    #                 ParsedElement(index=len(elements), kind="header",
    #                               text=text, section_path=tuple(heading_stack)))
    #     for footer in (section.footer, section.first_page_footer, section.even_page_footer):
    #         if footer is None:
    #             continue
    #         text = "\n".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
    #         if text:
    #             elements.append(
    #                 ParsedElement(index=len(elements), kind="footer",
    #                               text=text, section_path=tuple(heading_stack)))

    return elements, image_tasks


# ---------------------------------------------------------------------------
# Step 3 — resolve images via python-docx relationship API, describe, insert
# ---------------------------------------------------------------------------

_DESCRIBE_PROMPT = (
    "请详细描述这张图片的内容，包括图片中的文字、图表、流程图、表格等关键信息，用中文描述。"
)


def _describe_and_insert_images(
    doc: "Document",
    elements: list[ParsedElement],
    image_tasks: list[dict],
    describer: ImageDescriber,
) -> list[ParsedElement]:
    """Resolve image relationships, describe each image, insert descriptions.

    Uses python-docx's part.rels to look up relationship targets and read
    image blobs — no bespoke ZIP / XML handling.
    """
    if not image_tasks:
        return elements

    from collections import defaultdict

    rels = doc.part.rels
    insertions: dict[int, list[ParsedElement]] = defaultdict(list)

    for task in image_tasks:
        rel = rels.get(task["rel_id"])
        if rel is None:
            continue
        image_data = getattr(rel.target_part, "blob", None)
        if not image_data:
            continue

        # Describe — errors propagate up and abort.
        description = describer(image_data, _DESCRIBE_PROMPT)

        insert_at = task["paragraph_index"]
        insertions[insert_at].append(
            ParsedElement(
                index=-1,
                kind="image_description",
                text=description.strip(),
            )
        )

    if not insertions:
        return elements

    # Merge and re-index.
    result: list[ParsedElement] = []
    for i, el in enumerate(elements):
        result.append(el)
        if i in insertions:
            result.extend(insertions[i])
    for idx, el in enumerate(result):
        object.__setattr__(el, "index", idx)
    return result


# ---------------------------------------------------------------------------
# Shared helpers (unchanged from before except _heading_level)
# ---------------------------------------------------------------------------

def _sections_from_elements(elements: list[ParsedElement]) -> list[ParsedSection]:
    """Group DOCX elements into logical sections using heading styles."""
    sections: list[ParsedSection] = []
    current_title: str | None = None
    current_path: tuple[str, ...] = ()
    current_lines: list[str] = []

    for element in elements:
        if element.kind == "heading":
            if current_title is not None or current_lines:
                sections.append(
                    ParsedSection(
                        index=len(sections),
                        title=current_title,
                        path=current_path,
                        text="\n\n".join(current_lines).strip(),
                    )
                )
            current_title = element.text
            current_path = element.section_path
            current_lines = []
            continue
        current_lines.append(element.text)

    if current_title is not None or current_lines:
        sections.append(
            ParsedSection(
                index=len(sections),
                title=current_title,
                path=current_path,
                text="\n\n".join(current_lines).strip(),
            )
        )
    if sections:
        return sections
    return [ParsedSection(index=0, title="DOCX Document", text="")]


def _heading_level_from_style(style_name: str | None) -> int | None:
    """Map python-docx style name (e.g. 'Heading 1') to integer level."""
    if style_name is None:
        return None
    name = style_name.lower().replace(" ", "")
    if not name.startswith("heading"):
        return None
    suffix = name.removeprefix("heading")
    if not suffix.isdigit():
        return None
    level = int(suffix)
    return level if 1 <= level <= 6 else None